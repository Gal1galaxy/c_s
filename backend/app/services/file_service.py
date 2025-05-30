import os
from werkzeug.utils import secure_filename
from app import db
from app.models.file import File
from app.utils.crypto import AESCipher
from config import Config
import magic
import shutil
import tempfile
from datetime import datetime
import time
from flask import current_app
import pandas as pd
import base64
import re
from urllib.parse import unquote
from app.models.operation_log import OperationLog
import docx  # 需要安装 python-docx
from pdf2image import convert_from_path  # 需要安装 pdf2image
import fitz  # 需要安装 PyMuPDF
import io
from PIL import Image
import pyexcel as pe  # 用于读取 .xls 文件
import re



class FileService:
    def __init__(self):
        self.aes = AESCipher(Config.AES_KEY, Config.AES_IV)
        
    def secure_filename_with_chinese(self, filename):
        """安全的文件名处理，支持中文"""
        # URL 解码（处理可能的编码字符）
        filename = unquote(filename)
        
        # 移除路径分隔符和空白字符
        filename = re.sub(r'[\\/*?:"<>|]', '', filename)
        filename = filename.strip()
        
        # 如果文件名为空，返回默认名称
        if not filename:
            return 'untitled'
            
        return filename
        
    def log_operation(self, user_id, file_id, operation_type, operation_detail=None):
        """记录文件操作日志"""
        try:            
            # 检查文件是否存在
            file = File.query.get(file_id)
            if not file:
                print(f"File {file_id} not found, skipping log")
                return
            
            log = OperationLog(
                user_id=user_id,
                file_id=file_id,
                operation_type=operation_type,
                operation_detail=operation_detail
            )
            db.session.add(log)
            db.session.commit()

            print(log)
            
        except Exception as e:
            print(f"Error logging operation: {str(e)}")
            db.session.rollback()
        
    def save_file(self, file, user_id):
        """保存上传的文件"""
        try:
            # 使用支持中文的文件名处理
            filename = self.secure_filename_with_chinese(file.filename)
            file_extension = os.path.splitext(filename.lower())[1]
            
            # 根据文件扩展名设置用户友好的文件类型
            file_type = self.get_friendly_file_type(file_extension)
            
            # 创建用户的存储目录
            user_dir = os.path.join(Config.UPLOAD_FOLDER, str(user_id))
            if not os.path.exists(user_dir):
                os.makedirs(user_dir)
            
            # 检查文件名是否已存在（在数据库中）
            existing_file = File.query.filter_by(
                owner_id=user_id,
                filename=filename
            ).first()
            
            if existing_file:
                raise ValueError(f'已存在同名文件：{filename}')
            
            # 设置文件保存路径
            file_path = os.path.join(user_dir, filename)
            
            try:
                # 直接加密并保存文件对象
                encrypted_data = self.aes.encrypt_file(file)
                with open(file_path, 'wb') as f:
                    f.write(encrypted_data)
            except Exception as e:
                print(f"Error encrypting and saving file: {str(e)}")
                if os.path.exists(file_path):
                    os.remove(file_path)
                raise
            
            # 创建文件记录
            db_file = File(
                filename=filename,
                file_path=file_path,
                file_type=file_type,
                file_size=os.path.getsize(file_path),
                owner_id=int(user_id)
            )
            
            db.session.add(db_file)
            db.session.commit()
            
            # 记录上传操作
            self.log_operation(
                user_id=user_id,
                file_id=db_file.id,
                operation_type='upload',
                operation_detail=f'上传文件：{filename}'
            )
            
            return db_file
        except ValueError as e:
            # 文件名重复错误
            print(f"Duplicate filename error: {str(e)}")
            raise
        except Exception as e:
            print(f"Error in save_file: {str(e)}")
            db.session.rollback()
            raise
        
    def get_friendly_file_type(self, extension):
        """获取用户友好的文件类型显示"""
        type_map = {
            '.xlsx': 'Excel',
            '.xls': 'Excel',
            '.doc': 'Word',
            '.docx': 'Word',
            '.pdf': 'PDF',
            '.txt': '文本文件',
            '.jpg': '图片',
            '.jpeg': '图片',
            '.png': '图片',
            '.gif': '图片',
            '.mp4': '视频',
            '.mp3': '音频',
            '.zip': '压缩文件',
            '.rar': '压缩文件',
            '.7z': '压缩文件',
        }
        
        return type_map.get(extension.lower(), '其他文件')
        
    def get_decrypted_file_path(self, file):
        """获取解密后的临时文件路径"""
        try:
            print(f"Decrypting file: {file.filename}")  # 调试日志
            
            # 检查原始文件是否存在
            if not os.path.exists(file.file_path):
                raise FileNotFoundError(f"Original file not found: {file.file_path}")
            
            # 创建临时文件
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, file.filename)
            
            print(f"Reading encrypted file from: {file.file_path}")  # 调试日志
            
            # 读取加密文件
            with open(file.file_path, 'rb') as f:
                encrypted_data = f.read()
            
            print(f"Decrypting data of size: {len(encrypted_data)}")  # 调试日志
            
            # 解密文件
            decrypted_data = self.aes.decrypt_file(encrypted_data)
            
            print(f"Writing decrypted file to: {temp_path}")  # 调试日志
            
            # 保存解密后的临时文件
            with open(temp_path, 'wb') as f:
                f.write(decrypted_data)
            
            return temp_path
        except Exception as e:
            print(f"Error in get_decrypted_file_path: {str(e)}")  # 调试日志
            raise
        
    def delete_file(self, file):
        """删除文件"""
        try:
            # 保存文件信息用于日志记录
            file_info = {
                'id': file.id,
                'owner_id': file.owner_id,
                'filename': file.filename
            }
            
            # 开启事务
            db.session.begin()
            
            try:
                #db.session.commit()
                self.log_operation(
                    user_id=file_info['owner_id'],
                    file_id=file.id,
                    operation_type='delete',
                    operation_detail=f'删除文件：{file_info["filename"]}'
                )

                # 删除物理文件
                if os.path.exists(file.file_path):
                    os.remove(file.file_path)
                
                # 删除数据库记录
                db.session.delete(file)
                
                # 提交删除操作
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                print(f"Error in delete_file transaction: {str(e)}")
                raise
                
        except Exception as e:
            print(f"Error in delete_file: {str(e)}")
            db.session.rollback()
            raise
        
    def can_preview(self, filename):
        """检查文件是否支持预览"""
        extension = os.path.splitext(filename.lower())[1]
        previewable_extensions = {
            '.xlsx', '.xls',  # Excel文件
            '.docx', '.doc',  # Word文件
            '.pdf',          # PDF文件
            '.txt',          # 文本文件
            '.md'           # Markdown文件
        }
        return extension in previewable_extensions

    def to_dict(self, file):
        """将文件对象转换为字典"""
        return {
            'id': file.id,
            'filename': file.filename,
            'file_type': file.file_type,
            'file_size': file.file_size,
            'created_at': file.created_at.isoformat() if file.created_at else None,
            'updated_at': file.updated_at.isoformat() if file.updated_at else None,
            'is_public': file.is_public,
            'owner_id': file.owner_id,
            'can_preview': self.can_preview(file.filename)  # 添加预览支持标志
        }
        
    def update_file(self, original_file, new_file):
        """更新文件内容"""
        try:
            # 保存新文件
            new_file.save(original_file.file_path)
            
            # 更新文件信息
            original_file.file_size = os.path.getsize(original_file.file_path)
            original_file.updated_at = datetime.utcnow()
            
            db.session.commit()
            return original_file
        except Exception as e:
            db.session.rollback()
            raise e

    def get_file_content(self, file):
        """获取文件内容"""
        try:
            decrypted_path = self.get_decrypted_file_path(file)
            try:
                # Excel 文件处理
                if file.filename.lower().endswith(('.xlsx', '.xls')):
                    return self._handle_excel_file(decrypted_path)
                    
                # Word 文件处理
                elif file.filename.lower().endswith(('.docx', '.doc')):
                    return self._handle_word_file(decrypted_path)
                    
                # PDF 文件处理
                elif file.filename.lower().endswith('.pdf'):
                    return self._handle_pdf_file(decrypted_path)
                    
                # 文本文件处理
                elif file.filename.lower().endswith('.txt'):
                    return self._handle_txt_file(decrypted_path)
                    
                # 图片文件处理
                elif file.filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                    return self._handle_image_file(decrypted_path)
                    
                # ... 其他文件类型的处理保持不变 ...
                
            finally:
                if os.path.exists(decrypted_path):
                    os.remove(decrypted_path)
                    
        except Exception as e:
            print(f"Error in get_file_content: {str(e)}")
            raise
            
    def _handle_word_file(self, file_path):
        """处理 Word 文件"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            # 先处理所有段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # 只处理非空段落
                    content.append({
                        'type': 'paragraph',
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # 再处理所有表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # 获取单元格中的所有段落的文本
                        cell_text = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                        row_data.append(cell_text)
                    if any(row_data):  # 只添加非空行
                        table_data.append(row_data)
                
                if table_data:  # 只添加非空表格
                    content.append({
                        'type': 'table',
                        'data': table_data
                    })
            
            return {
                'content': content,
                'file_type': 'Word'
            }
            
        except Exception as e:
            print(f"Error processing Word file: {str(e)}")
            raise
            
    def _handle_pdf_file(self, file_path):
        """处理 PDF 文件"""
        try:
            # 使用 PyMuPDF 打开 PDF
            pdf_document = fitz.open(file_path)
            content = []
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # 将页面渲染为图片
                zoom = 2  # 设置缩放比例以提高图片质量
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                
                # 将图片转换为 base64
                img_data = pix.tobytes("png")  # 直接输出为 PNG 格式
                img_base64 = base64.b64encode(img_data).decode()
                
                content.append({
                    'page': page_num + 1,
                    'image': img_base64,
                    'width': pix.width,
                    'height': pix.height
                })
                
            pdf_document.close()
            
            return {
                'content': content,
                'file_type': 'PDF',
                'total_pages': len(content)
            }
            
        except Exception as e:
            print(f"Error processing PDF file: {str(e)}")
            raise

    def _handle_txt_file(self, file_path):
        """处理 txt 文件"""
        try:
            # 使用 PyMuPDF 打开 PDF
            with open(file_path, 'r') as f:
                content = f.read()
            
            return {
                'content': content,
                'file_type': 'text/plain',
                'total_pages': len(content)
            }
            
        except Exception as e:
            print(f"Error processing PDF file: {str(e)}")
            raise

    def _handle_image_file(self, file_path):
        """处理图片文件"""
        try:
            # 打开图片
            with Image.open(file_path) as image:
                # 创建一个字节缓冲区
                buffer = io.BytesIO()
                
                # 保存图片到缓冲区，格式为PNG
                image.save(buffer, format='PNG')
                
                # 获取字节数据并转换为base64
                image_base64 = base64.b64encode(buffer.getvalue()).decode()
                
                return {
                    'content': image_base64,
                    'file_type': 'image/png',  # 统一使用PNG格式
                    'total_pages': 1,
                    'width': image.width,
                    'height': image.height
                }
                
        except Exception as e:
            print(f"Error processing image file: {str(e)}")
            raise

    def _handle_excel_file(self, file_path):
        """处理 Excel 文件（支持 .xlsx 和 .xls，无伪表头判断）"""
        try:
            content = {}

            # 优先尝试 openpyxl 读取 .xlsx
            try:
                df_dict = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
                for sheet_name, df in df_dict.items():
                    print(f"📄 正在处理 sheet: {sheet_name}")
                    if df.empty:
                        headers = [str(h).strip() for h in df.columns.tolist()]
                        print(f"✅ 读取空表的表头（xlsx）: {headers}")
                        sheet_content = [
                            {str(i): h for i, h in enumerate(headers)}
                        ]
                        content[sheet_name] = sheet_content
                        continue

                    headers = [str(h).strip() for h in df.columns.tolist()]
                    print(f"✅ 读取表头（xlsx）: {headers}")
                    sheet_content = [
                        {str(i): h for i, h in enumerate(headers)}
                    ] + [
                        {str(i): str(row[h]) if pd.notna(row[h]) else '' for i, h in enumerate(headers)}
                        for _, row in df.iterrows()
                    ]
                    content[sheet_name] = sheet_content

                return {
                    'content': content,
                    'file_type': 'Excel'
                }

            except Exception as openpyxl_error:
                print(f"[Fallback] openpyxl failed: {openpyxl_error}")

                try:
                    book = pe.get_book(file_name=file_path)
                    for sheet in book:
                        print(f"📄 正在处理 sheet: {sheet.name}")
                        rows = sheet.to_array()
                        if not rows:
                            content[sheet.name] = []
                            continue

                        headers = [str(h).strip() for h in rows[0]]
                        print(f"✅ 读取表头（xls）: {headers}")
                        sheet_content = [
                            {str(i): h for i, h in enumerate(headers)}
                        ] + [
                            {str(i): str(cell) if cell is not None else '' for i, cell in enumerate(row)}
                            for row in rows[1:]
                        ]
                        content[sheet.name] = sheet_content

                    return {
                        'content': content,
                        'file_type': 'Excel'
                    }

                except Exception as pe_error:
                    print(f"[Error] Both openpyxl and pyexcel failed.")
                    raise pe_error
    
        except Exception as e:
            print(f"Error processing Excel file: {str(e)}")
            raise



    def update_file_content(self, file, content):
        """更新文件内容"""
        temp_path = None
        try:
            temp_path = os.path.join(current_app.config['TEMP_FOLDER'], f'temp_{file.id}_{int(time.time())}')
            print(f"Updating content for file: {file.filename}")

            if file.file_type.endswith('spreadsheet') or file.filename.lower().endswith(('.xlsx', '.xls')):
                try:
                    print(f"Processing Excel content: {content}")
    
                    writer = pd.ExcelWriter(temp_path, engine='openpyxl')  # 明确指定 openpyxl
                    for sheet_name, sheet_data in content.items():
                        print(f"Sheet: {sheet_name}: {sheet_data}")

                        if not sheet_data:
                            print(f"[警告] Sheet {sheet_name} 内容为空，已跳过写入")
                            continue

                        header_row = sheet_data[0]
                        data = sheet_data[1:]

                        if not isinstance(header_row, dict):
                            print(f"[错误] 表头格式不合法（不是 dict），跳过写入 Sheet {sheet_name}")
                            continue

                        header_keys = list(header_row.keys())
                        seen = set()
                        header_names = []
                        valid_header_count = 0

                        for i, k in enumerate(header_keys):
                            raw = header_row[k].strip() if isinstance(header_row[k], str) else str(header_row[k]).strip()
                            col_name = raw if raw else f"列{i}"

                            # 记录是否为有效列名（非占位符）
                            if not re.match(r'^(列\d+|Unnamed.*|\d+|\s*)$', col_name):
                                valid_header_count += 1

                            while col_name in seen:
                                col_name += '_1'
                            seen.add(col_name)
                            header_names.append(col_name)

                        # ✅ 拦截非法表头：如果有效字段数不足 2，跳过写入
                        if valid_header_count < 2:
                            print(f"[警告] Sheet {sheet_name} 表头无效（有效字段数为 {valid_header_count}），跳过写入")
                            continue

                        print("✅ 过滤后表头:", header_names)

                        rows = [
                            [row.get(k, '') for k in header_keys]
                            for row in data
                        ]

                        df = pd.DataFrame(rows, columns=header_names)
                        print("✅ DataFrame:\n", df.head())

                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                        print(f"✅ Written sheet: {sheet_name} with columns: {df.columns}")
    
                    writer.close()

                    with open(temp_path, 'rb') as f:
                        file_data = f.read()
                    print(f"Excel file size before encryption: {len(file_data)}")
    
                except Exception as e:
                    print(f"Error processing Excel file: {str(e)}")
                    raise
            else:
                return
    
            # 加密保存
            encrypted_data = self.aes.encrypt_file(file_data)
            with open(file.file_path, 'wb') as f:
                f.write(encrypted_data)
    
            file.file_size = os.path.getsize(file.file_path)
            file.updated_at = datetime.utcnow()
            #储存为.xlsx,并更新文件名
            if file.filename.lower().endswith('.xls'):
                new_filename = file.filename[:-4] + '.xlsx'
                new_path = os.path.join(os.path.dirname(file.file_path), new_filename)

                os.rename(file.file_path, new_path)  # 重命名加密后文件
                file.file_path = new_path
                file.filename = new_filename

            #统一设置文件类型为openxml格式
            file.file_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            db.session.commit()
            print(f"File updated successfully: {file.filename}")

            self.log_operation(
                user_id=file.owner_id,
                file_id=file.id,
                operation_type='edit',
                operation_detail=f'编辑文件:{file.filename}'
            )

            return True

        except Exception as e:
            print(f"Error in update_file_content: {str(e)}")
            db.session.rollback()
            raise
        finally:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)
